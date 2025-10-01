"""
Utilidades para extracción de texto de diferentes formatos de archivo
"""
import os
from typing import Tuple
from PyPDF2 import PdfReader
from docx import Document
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class TextExtractor:
    """Extrae texto de documentos en diferentes formatos"""

    SUPPORTED_FORMATS = {'.pdf', '.txt', '.docx'}

    @staticmethod
    def extract_text(file_path: str, filename: str) -> Tuple[str, str]:
        """
        Extrae texto de un archivo

        Args:
            file_path: Ruta al archivo temporal
            filename: Nombre original del archivo

        Returns:
            Tuple de (texto_extraído, tipo_documento)
        """
        _, ext = os.path.splitext(filename.lower())

        if ext not in TextExtractor.SUPPORTED_FORMATS:
            raise ValueError(f"Formato no soportado: {ext}. Formatos permitidos: {TextExtractor.SUPPORTED_FORMATS}")

        try:
            if ext == '.pdf':
                return TextExtractor._extract_from_pdf(file_path), 'pdf'
            elif ext == '.txt':
                return TextExtractor._extract_from_txt(file_path), 'txt'
            elif ext == '.docx':
                return TextExtractor._extract_from_docx(file_path), 'docx'
        except Exception as e:
            logger.error(f"Error extrayendo texto de {filename}: {str(e)}")
            raise

    @staticmethod
    def _extract_from_pdf(file_path: str) -> str:
        """Extrae texto de un PDF"""
        reader = PdfReader(file_path)
        text = []

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)

        full_text = '\n\n'.join(text)
        logger.info(f"Extraídas {len(reader.pages)} páginas del PDF")
        return full_text

    @staticmethod
    def _extract_from_txt(file_path: str) -> str:
        """Extrae texto de un archivo TXT"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()

        logger.info(f"Extraído texto de archivo TXT ({len(text)} caracteres)")
        return text

    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        """Extrae texto de un documento DOCX"""
        doc = Document(file_path)
        text = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)

        full_text = '\n\n'.join(text)
        logger.info(f"Extraídos {len(doc.paragraphs)} párrafos del DOCX")
        return full_text

    @staticmethod
    def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
        """
        Divide texto largo en chunks con overlap

        Args:
            text: Texto completo a dividir
            chunk_size: Tamaño de cada chunk en tokens (aproximado)
            overlap: Cantidad de tokens de overlap entre chunks

        Returns:
            Lista de chunks de texto
        """
        # Aproximación simple: 1 token ≈ 4 caracteres
        char_chunk_size = chunk_size * 4
        char_overlap = overlap * 4

        chunks = []
        start = 0

        while start < len(text):
            end = start + char_chunk_size
            chunk = text[start:end]

            # Intentar cortar en límite de palabra
            if end < len(text):
                last_space = chunk.rfind(' ')
                if last_space > char_chunk_size * 0.8:  # Si hay un espacio en el último 20%
                    chunk = chunk[:last_space]
                    end = start + last_space

            if chunk.strip():
                chunks.append(chunk.strip())

            start = end - char_overlap

        logger.info(f"Texto dividido en {len(chunks)} chunks")
        return chunks
